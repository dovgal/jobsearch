from __future__ import annotations

import json
import threading
import uuid
from pathlib import Path
from typing import Any, Optional

import io

import yaml
try:
    import docx as python_docx
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
from fastapi import FastAPI, Form, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from core.session import Session
from core.settings import Settings
from pipeline.orchestrator import Orchestrator

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent
REPO_DIR = BASE_DIR.parent
OUTPUT_DIR = REPO_DIR / "output"
CONFIG_PATH = REPO_DIR / "config" / "config.yaml"
ENV_PATH = REPO_DIR / ".env"
OUTPUT_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# .env helpers
# ---------------------------------------------------------------------------

def _read_env() -> dict[str, str]:
    result: dict[str, str] = {}
    if not ENV_PATH.exists():
        return result
    for line in ENV_PATH.read_text("utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "=" in line:
            k, _, v = line.partition("=")
            result[k.strip()] = v.strip()
    return result


def _write_env(updates: dict[str, str]) -> None:
    lines: list[str] = []
    existing_keys: set[str] = set()

    if ENV_PATH.exists():
        for line in ENV_PATH.read_text("utf-8").splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith("#"):
                lines.append(line)
                continue
            if "=" in stripped:
                k = stripped.partition("=")[0].strip()
                existing_keys.add(k)
                if k in updates:
                    lines.append(f"{k}={updates[k]}")
                else:
                    lines.append(line)
            else:
                lines.append(line)

    for k, v in updates.items():
        if k not in existing_keys:
            lines.append(f"{k}={v}")

    ENV_PATH.write_text("\n".join(lines) + "\n", "utf-8")

# ---------------------------------------------------------------------------
# Word (.docx) helpers
# ---------------------------------------------------------------------------

def _extract_docx_text(data: bytes) -> str:
    """Извлечь plain-text из .docx файла."""
    if not _DOCX_AVAILABLE:
        return ""
    doc = Document(io.BytesIO(data))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def _apply_run_style(run, bold: bool = False, font_name: str = "Open Sans",
                     font_size_pt: float | None = None,
                     color_hex: str = "333333") -> None:
    run.bold = bold
    run.font.name = font_name
    if font_size_pt:
        run.font.size = python_docx.shared.Pt(font_size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_styled_para(doc: Document, text: str, style: str,
                     bold: bool = False, font_size_pt: float | None = None,
                     color_hex: str = "333333", align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    """Добавить параграф с явным применением шрифта Open Sans."""
    parts = text.split("**")
    p = doc.add_paragraph(style=style)
    p.alignment = align
    for i, part in enumerate(parts):
        if not part:
            continue
        is_bold = (i % 2 == 1) or bold
        run = p.add_run(part)
        _apply_run_style(run, bold=is_bold, font_size_pt=font_size_pt, color_hex=color_hex)


def _markdown_to_docx(md_text: str, template_path: str | Path | None = None) -> bytes:
    """Конвертировать markdown в .docx, сохраняя стиль шаблона если задан."""

    # Открываем шаблон (оригинальный CV) либо чистый документ
    if template_path and Path(template_path).exists():
        doc = Document(str(template_path))
        # Очищаем всё содержимое, сохраняя стили и поля
        from docx.oxml.ns import qn as _qn
        body = doc.element.body
        sect_pr = body.find(_qn("w:sectPr"))
        for child in list(body):
            if child is not sect_pr:
                body.remove(child)
    else:
        doc = Document()
        section = doc.sections[0]
        section.page_width = python_docx.shared.Cm(21)
        section.page_height = python_docx.shared.Cm(29.7)
        section.top_margin = python_docx.shared.Cm(2)
        section.bottom_margin = python_docx.shared.Cm(0.5)
        section.left_margin = python_docx.shared.Cm(2.5)
        section.right_margin = python_docx.shared.Cm(2.5)

    # Определяем доступные стили (из шаблона или стандартные)
    style_names = {s.name for s in doc.styles}
    S_TITLE   = "Title"       if "Title" in style_names else "Normal"
    S_NORMAL  = "Normal"
    S_BODY    = "Body Text"   if "Body Text" in style_names else "Normal"
    S_BULLET  = "Body Text"   if "Body Text" in style_names else "List Bullet"

    for line in md_text.splitlines():
        stripped = line.rstrip()

        # # Имя / заголовок первого уровня
        if stripped.startswith("# "):
            _add_styled_para(doc, stripped[2:], S_TITLE, bold=True,
                             font_size_pt=16, color_hex="000000")

        # ## Секция (PROFIL, EXPÉRIENCES, etc.)
        elif stripped.startswith("## "):
            _add_styled_para(doc, stripped[3:].upper(), S_TITLE, bold=True,
                             font_size_pt=12, color_hex="333333")

        # ### Должность / дата (bold 10pt)
        elif stripped.startswith("### "):
            _add_styled_para(doc, stripped[4:], S_NORMAL, bold=True,
                             font_size_pt=10, color_hex="333333")

        # #### Компания (bold 10pt)
        elif stripped.startswith("#### "):
            _add_styled_para(doc, stripped[5:], S_NORMAL, bold=True,
                             font_size_pt=10, color_hex="333333")

        # - bullet
        elif stripped.startswith(("- ", "* ")):
            _add_styled_para(doc, stripped[2:], S_BULLET, bold=False,
                             color_hex="333333")

        # Разделитель
        elif stripped == "---" or stripped == "":
            doc.add_paragraph("", style=S_BODY)

        # Обычный текст / контактная информация
        else:
            _add_styled_para(doc, stripped, S_BODY, bold=False,
                             color_hex="333333")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------

app = FastAPI(title="Job Search Consultation", docs_url=None, redoc_url=None)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

# ---------------------------------------------------------------------------
# Background task registry
# ---------------------------------------------------------------------------

_tasks: dict[str, dict[str, Any]] = {}
_lock = threading.Lock()


def _set_task(task_id: str, data: dict) -> None:
    with _lock:
        _tasks[task_id] = data


def _get_task(task_id: str) -> dict:
    with _lock:
        return _tasks.get(task_id, {"status": "unknown"})


def _start(task_id: str, fn, *args, **kwargs) -> None:
    _set_task(task_id, {"status": "running"})

    def _run():
        try:
            result = fn(*args, **kwargs)
            _set_task(task_id, {"status": "done", "result": result})
        except Exception as exc:
            _set_task(task_id, {"status": "error", "error": str(exc)})

    threading.Thread(target=_run, daemon=True).start()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _session(sid: str) -> Session:
    return Session.open(sid, base_dir=str(OUTPUT_DIR))


def _cv_path(sid: str) -> Path:
    return OUTPUT_DIR / sid / "input_cv.txt"


def _session_ctx(sid: str) -> dict:
    """Build template context from session state."""
    session = _session(sid)
    state = session.state
    cv_text = _cv_path(sid).read_text("utf-8") if _cv_path(sid).exists() else ""

    def _exists(key: str) -> bool:
        p = state.get(key)
        return bool(p and Path(p).exists())

    questions_path = OUTPUT_DIR / sid / "strategy" / "coaching_questions.json"
    questions = (
        json.loads(questions_path.read_text("utf-8"))
        if questions_path.exists()
        else []
    )
    intake_path = OUTPUT_DIR / sid / "strategy" / "intake_answers.json"
    has_intake = intake_path.exists()

    jobs: list[dict] = []
    if _exists("jobs_path"):
        jobs = json.loads(Path(state["jobs_path"]).read_text("utf-8"))

    return {
        "sid": sid,
        "cv_preview": cv_text[:400] + ("…" if len(cv_text) > 400 else ""),
        "state": state,
        "has_strategy": _exists("strategy_path"),
        "has_base_cv": _exists("base_cv_path"),
        "has_linkedin": _exists("linkedin_path"),
        "has_jobs": _exists("jobs_path"),
        "questions": questions,
        "has_intake": has_intake,
        "jobs": [j for j in jobs if (j.get("tier") or "skip") != "skip"][:30],
    }


# ---------------------------------------------------------------------------
# Routes — pages
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    sessions = []
    for p in sorted(OUTPUT_DIR.iterdir(), key=lambda x: x.stat().st_mtime, reverse=True):
        if not p.is_dir():
            continue
        state = {}
        state_file = p / "state.json"
        if state_file.exists():
            state = json.loads(state_file.read_text("utf-8"))
        sessions.append({
            "id": p.name,
            "has_strategy": bool(state.get("strategy_path")),
            "has_jobs": bool(state.get("jobs_path")),
        })
    return templates.TemplateResponse(request, "home.html", {"sessions": sessions[:10]})


@app.post("/session/new")
async def new_session(cv_text: str = Form(""), cv_file: Optional[UploadFile] = None):
    raw_docx: bytes | None = None
    if cv_file and cv_file.filename:
        raw = await cv_file.read()
        if cv_file.filename.lower().endswith(".docx"):
            content = _extract_docx_text(raw).strip()
            raw_docx = raw
        else:
            content = raw.decode("utf-8", errors="ignore").strip()
    else:
        content = cv_text.strip()

    if not content:
        return RedirectResponse("/", status_code=303)

    session = Session.open(base_dir=str(OUTPUT_DIR))
    sid = session.session_id
    _cv_path(sid).write_text(content, "utf-8")
    # Сохраняем оригинальный .docx как шаблон стиля
    if raw_docx:
        (OUTPUT_DIR / sid / "input_cv.docx").write_bytes(raw_docx)
    return RedirectResponse(f"/session/{sid}", status_code=303)


@app.get("/session/{sid}", response_class=HTMLResponse)
async def session_page(request: Request, sid: str):
    ctx = _session_ctx(sid)
    return templates.TemplateResponse(request, "session.html", ctx)


# ---------------------------------------------------------------------------
# Routes — pipeline partial (refreshed after each task)
# ---------------------------------------------------------------------------

@app.get("/session/{sid}/pipeline", response_class=HTMLResponse)
async def pipeline_partial(request: Request, sid: str):
    ctx = _session_ctx(sid)
    return templates.TemplateResponse(request, "partials/pipeline.html", ctx)


# ---------------------------------------------------------------------------
# Routes — start tasks
# ---------------------------------------------------------------------------

@app.post("/session/{sid}/run/questions", response_class=HTMLResponse)
async def run_questions(request: Request, sid: str):
    task_id = uuid.uuid4().hex[:8]

    def _work():
        from agents.strategy_agent import StrategyAgent
        s = _session(sid)
        cv = _cv_path(sid).read_text("utf-8")
        agent = StrategyAgent(Settings.load(), s)
        qs = agent.generate_questions(cv)
        s.write_json("strategy/coaching_questions.json", qs)
        return {"questions": qs}

    _start(task_id, _work)
    return templates.TemplateResponse(request, "partials/task_running.html", {
        "task_id": task_id, "sid": sid,
        "message": "Анализирую CV, формулирую вопросы…",
    })


@app.post("/session/{sid}/intake", response_class=HTMLResponse)
async def submit_intake(request: Request, sid: str, answers_json: str = Form("")):
    """User submitted coaching answers → run full strategy."""
    task_id = uuid.uuid4().hex[:8]
    answers = json.loads(answers_json) if answers_json.strip() else {}

    def _work():
        s = _session(sid)
        intake_path = s.write_json("strategy/intake_answers.json", answers)
        orc = Orchestrator(session=s, cv_path=str(_cv_path(sid)))
        return orc.strategy(intake_path=str(intake_path))

    _start(task_id, _work)
    return templates.TemplateResponse(request, "partials/task_running.html", {
        "task_id": task_id, "sid": sid,
        "message": "Провожу веб-ресёрч и пишу стратегию (~5 мин)…",
    })


@app.post("/session/{sid}/run/{agent}", response_class=HTMLResponse)
async def run_agent(request: Request, sid: str, agent: str,
                    job_id: str = Form("")):
    task_id = uuid.uuid4().hex[:8]
    labels = {
        "base-cv":    "Переписываю CV под стратегическое позиционирование…",
        "linkedin":   "Создаю LinkedIn-профиль…",
        "hunt":       "Сканирую каналы вакансий (~5–10 мин)…",
        "bulk-resume":"Пакетная адаптация Tier B…",
        "content":    "Пишу контент-план и посты…",
        "interview":  "Готовлю пакет к интервью…",
        "resume":     f"Адаптирую резюме под вакансию {job_id}…",
        "outreach":   f"Готовлю outreach для вакансии {job_id}…",
    }

    def _work():
        s = _session(sid)
        orc = Orchestrator(session=s, cv_path=str(_cv_path(sid)))
        dispatch = {
            "base-cv":    orc.base_cv,
            "linkedin":   orc.linkedin,
            "hunt":       orc.hunt,
            "bulk-resume": lambda: orc.bulk_resume("B"),
            "content":    orc.content,
            "interview":  lambda: orc.interview(job_id or None),
            "resume":     lambda: orc.resume(job_id),
            "outreach":   lambda: orc.outreach(job_id),
        }
        return dispatch[agent]()

    _start(task_id, _work)
    return templates.TemplateResponse(request, "partials/task_running.html", {
        "task_id": task_id, "sid": sid,
        "message": labels.get(agent, "Запускаю…"),
    })


# ---------------------------------------------------------------------------
# Routes — task polling
# ---------------------------------------------------------------------------

@app.get("/session/{sid}/task/{task_id}", response_class=HTMLResponse)
async def task_status(request: Request, sid: str, task_id: str):
    task = _get_task(task_id)
    status = task["status"]

    if status == "running":
        return templates.TemplateResponse(request, "partials/task_running.html", {
            "task_id": task_id, "sid": sid,
            "message": "Выполняется…",
        })

    if status == "error":
        return templates.TemplateResponse(request, "partials/task_error.html", {
            "sid": sid,
            "error": task.get("error", "Неизвестная ошибка"),
        })

    # Done — check if strategy returned questions
    result = task.get("result", {})
    if result.get("stage") == "awaiting_intake":
        ctx = _session_ctx(sid)
        return templates.TemplateResponse(request, "partials/questions.html", ctx)

    return templates.TemplateResponse(request, "partials/task_done.html", {
        "sid": sid, "result": result,
    })


# ---------------------------------------------------------------------------
# Routes — files
# ---------------------------------------------------------------------------

@app.get("/session/{sid}/vacancies", response_class=HTMLResponse)
async def vacancies_page(request: Request, sid: str):
    s = _session(sid)
    p = s.recall("vacancies_html")
    if p and Path(p).exists():
        return HTMLResponse(Path(p).read_text("utf-8"))
    return HTMLResponse("<p style='color:white;padding:2rem'>Запустите hunt сначала.</p>", 404)


@app.get("/session/{sid}/file/{fp:path}")
async def download_file(sid: str, fp: str):
    path = OUTPUT_DIR / sid / fp
    if path.is_file():
        return FileResponse(str(path), filename=path.name)
    return HTMLResponse("Файл не найден", 404)


@app.get("/session/{sid}/docx/{fp:path}")
async def download_docx(sid: str, fp: str):
    """Конвертировать .md файл в .docx, используя оригинальный CV как шаблон стиля."""
    if not _DOCX_AVAILABLE:
        return HTMLResponse("python-docx не установлен на сервере", 503)
    path = OUTPUT_DIR / sid / fp
    if not path.is_file():
        return HTMLResponse("Файл не найден", 404)
    md_text = path.read_text("utf-8")
    template = OUTPUT_DIR / sid / "input_cv.docx"
    docx_bytes = _markdown_to_docx(md_text, template_path=template if template.exists() else None)
    filename = path.stem + ".docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ---------------------------------------------------------------------------
# Settings
# ---------------------------------------------------------------------------

def _settings_ctx() -> dict:
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f) or {}

    env_vals = _read_env()
    providers_info = []
    for name, pcfg in (cfg.get("providers") or {}).items():
        key_env = pcfg.get("api_key_env", "")
        current_key = env_vals.get(key_env, "") if key_env else ""
        providers_info.append({
            "name": name,
            "label": name.capitalize(),
            "key_env": key_env,
            "model": pcfg.get("model", ""),
            "has_key": bool(current_key),
            "key_masked": ("●" * 8 + current_key[-4:]) if len(current_key) > 4 else ("●" * len(current_key) if current_key else ""),
        })

    return {
        "providers": providers_info,
        "default_provider": cfg.get("default_provider", "ollama"),
        "all_provider_names": [p["name"] for p in providers_info],
        "saved": False,
    }


@app.get("/settings", response_class=HTMLResponse)
async def settings_page(request: Request):
    ctx = _settings_ctx()
    return templates.TemplateResponse(request, "settings.html", ctx)


@app.post("/settings", response_class=HTMLResponse)
async def settings_save(request: Request):
    form = await request.form()

    # Update default_provider in config.yaml
    new_default = form.get("default_provider", "")
    if new_default:
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
        cfg["default_provider"] = new_default
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            yaml.dump(cfg, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

    # Update API keys in .env
    env_updates: dict[str, str] = {}
    for key, val in form.items():
        if key.startswith("key_") and val.strip():
            env_var = key[4:]  # strip "key_" prefix
            env_updates[env_var] = val.strip()

    if env_updates:
        _write_env(env_updates)

    ctx = _settings_ctx()
    ctx["saved"] = True
    return templates.TemplateResponse(request, "settings.html", ctx)


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@app.get("/health")
async def health():
    return {"status": "ok"}
